"""Module to create docs (student and teacher docs) to evaluate summary statistics.

Generates distributions with unique means and random standard deviations (std). The 
summary statistics evaluated are mean, median, mode, variance, standard deviation (std),
min value, and max value.

The mean may not be unique between exercises if:
    * conf["DISTRIBUTION_DISTANCE"] > conf["INTERVALS_MEAN"]['step']

If you want 15 problems in docs
    * conf['STATISTICS_M1']['N_QUESTION_TABLES'] must be 15.
If you want 20 samples per distribution:
    * conf['STATISTICS_M1']['DIST_SIZE'] must be 20.

    Typical usage example:
    from writers.statistics_m1 import StatisticsModule1
    <<loads config file in variable conf>>

    writer = StatisticsModule1(conf)
    writer.write()
"""

from collections import namedtuple, OrderedDict
import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
from scipy import stats as st
from scipy.stats import truncnorm

from writers.writer_interface import DocumentWriterInterface


class StatisticsModule1(DocumentWriterInterface):
    """Creates module_1 documents (student and teacher docs) for summary statistics."""

    def __init__(self, conf: dict) -> None:
        self.doc = docx.Document()
        self.conf = conf["STATISTICS_M1"]
        self.sol_table_name = self.conf["SOLUTIONS_TABLE_NAME"]
        self.solutions_path = self.conf["SOLUTIONS_PATH"]
        self.questions_path = self.conf["QUESTIONS_PATH"]
        self.n_question_tables = self.conf["N_QUESTION_TABLES"]
        self.heading_text = self.conf["HEADING_TEXT"]
        self.question_text = self.conf["QUESTION_TEXT"]
        self.footer_text = self.conf["FOOTER_TEXT"]
        self.intervals_mean = self.conf["INTERVALS_MEAN"]
        self.intervals_std = self.conf["INTERVALS_STD"]
        self.distribution_distance = self.conf["DISTRIBUTION_DISTANCE"]
        self.size = self.conf["DIST_SIZE"]

        self.Dist = namedtuple("Distribution", "values target_mean target_std")
        self.distributions = []
        self.answers = []

    def compute_questions_and_answers(self) -> None:
        self._create_distributions()
        self._create_answers()
        print(
            f"\t{len(self.distributions)} new distributions were created with their solutions."
        )

    def _create_distributions(self) -> None:
        """Generates `n_question_tables` arrays with unique means and random std."""
        max_i = (
            self.intervals_mean["min"]
            + self.intervals_mean["step"] * self.n_question_tables
        )
        for mean in range(
            self.intervals_mean["min"], max_i, self.intervals_mean["step"]
        ):
            std = np.random.randint(
                self.intervals_std["min"], self.intervals_std["max"], 1
            )[0]
            dist_tup = self._create_dist_with_tolerance(
                mean, std, self.distribution_distance, self.size
            )
            self.distributions.append(dist_tup)

    def _create_dist_with_tolerance(
        self,
        target_mean: int,
        target_std: int,
        max_distance: float = 2.0,
        size: int = 30,
    ) -> namedtuple:
        """Creates a random distribution near a mean target.

        Args:
            target_mean (int): mean target
            target_std (int): std target. The real std may be far from the target value
            max_distance (float, optional): max diff with target mean. Defaults to 2.0
            size (int, optional): number of samples in the distribution. Defaults to 30

        Returns:
            namedtuple: distribution values, target mean and target std
        """
        count = 0
        dist = self._create_int_normal_dist(target_mean, target_std, size)
        mean = dist.mean()
        distance = abs(target_mean - mean)
        best_distance, alternative_dist = distance, dist
        while distance >= max_distance and count < 30:
            dist = self._create_int_normal_dist(target_mean, target_std, size)
            mean = dist.mean()
            distance = abs(target_mean - mean)
            if best_distance > distance:
                best_distance, alternative_dist = distance, dist
            count += 1
        if count == 30:
            print("\tHard to find distribution. Giving the best option")
            dist = alternative_dist
        return self.Dist(dist, target_mean, target_std)

    def _create_int_normal_dist(self, mean: float, std: float, size: int) -> np.array:
        """Creates a normal-like random distribution of integers.

        Args:
            mean (float): target mean
            std (float): target standard deviation
            size (int): number of samples in the distribution

        Returns:
            np.array: dist with a mean and std maybe similar to `std` and `mean`.
        """
        dist = truncnorm(a=-6, b=6, scale=1).rvs(size=size)
        dist = dist.round().astype(int)
        return std * dist + mean

    def _create_answers(self) -> None:
        """Creates the answers list"""
        self.answers = [
            self._create_one_answer(dist.values) for dist in self.distributions
        ]

    def _create_one_answer(self, x: np.array) -> OrderedDict:
        """Computes summary statistics for a given array.

        For the mode, if there is more than one, retrieves just the smallest.

        Args:
            x (np.array): distribution of integers

        Returns:
            OrderedDict: returns 7 summary statistics
        """
        var = np.round(x.var(), 2)
        return OrderedDict(
            {
                "mean": np.round(x.mean(), 2),
                "median": int(np.median(x)),
                "mode": st.mode(x)[0][0],
                "var": var,
                "std": np.round(np.sqrt(var), 2),
                "x_min": x.min(),
                "x_max": x.max(),
            }
        )

    def create_questions_document(self) -> None:
        self.doc = docx.Document()
        all_dist_values = [dist.values for dist in self.distributions]
        heading = self.doc.add_heading(self.heading_text, 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = heading.style
        title_style.font.name = "Open Sans"

        for i in range(self.n_question_tables):
            para = self.doc.add_paragraph()
            for fragment in re.split("<b>|<i>", self.question_text):
                if fragment[-2:] == "*b":
                    para.add_run(fragment[:-2]).bold = True
                elif fragment[-2:] == "*i":
                    para.add_run(fragment[:-2]).italic = True
                else:
                    para.add_run(fragment)

            table = self.doc.add_table(rows=1, cols=self.size + 1)
            table.style = "Table Grid"
            row = table.rows[0]
            for j, cell in enumerate(row.cells):
                if j == 0:
                    cell.text = f"X_{i}"
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True
                else:
                    value = all_dist_values[i][j - 1]
                    cell.text = str(value)

            para = self.doc.add_paragraph(" - " * 30)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = self.doc.sections[0].footer
        para = footer.paragraphs[0]
        para.add_run(self.footer_text).italic = True
        self.doc.save(self.questions_path)
        print(f"\tQuestions saved in: {self.questions_path}")

    def create_solutions_document(self) -> None:
        self.doc = docx.Document()
        for i, answer in enumerate(self.answers):
            dist = self.distributions[i]
            name = f"{self.sol_table_name} {i} - target_mean {dist.target_mean} - target_std {dist.target_std}"
            self._write_one_answer_table(dist, answer, name)
        self.doc.save(self.solutions_path)
        print(f"\tAnswers saved in: {self.solutions_path}")

    def _write_one_answer_table(
        self, dist: namedtuple, answer: OrderedDict, name: str
    ) -> None:
        """Creates a single answer table.

        Args:
            dist (namedtuple): a distribution expressed as a self.Dist namedtuple
            answer (OrderedDict): a group of summary statistics
            name (str): title fo the table
        """
        p1 = self.doc.add_paragraph()
        p1.add_run(name).italic = True
        table = self.doc.add_table(rows=2, cols=len(answer))
        table.style = "Table Grid"
        keys = list(answer.keys())
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0:
                    cell.text = keys[j]
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True
                else:
                    cell.text = str(answer[keys[j]])
